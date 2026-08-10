import re
from io import BytesIO

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor
from docx.table import Table, _Cell

from app.agents.lesson_generate.schema import LessonOutput, Worksheet

# HTML(result.html)의 worksheet-section 색상과 맞춘 팔레트(교수학습과정안 표에서 사용).
_SECTION_COLORS = ["DCEBFB", "DFF5E6", "F8E4FC", "FDF0D5"]
_MISSION_COLOR = "FFF6E0"
_HEADER_COLOR = "D4E0EC"

# 활동지는 실제로 인쇄·복사해서 나눠주는 자료라 배경색을 꽉 채우면 잉크 낭비에다
# 흑백 복사 시 안 예쁘다. 참고 이미지(초등학교가 궁금해요/독서 기록장)처럼 흰 배경 +
# 테두리 박스 + 옅은 회색 라벨 줄로만 구성한다.
_WORKSHEET_LABEL_GRAY = "EFEFEF"
_WORKSHEET_FONT_NAME = "나눔바른펜"

# 원본 교수학습과정안 PDF에서 "학습 문제 안내하기" 행은 교사/학생 칸을 가로로 합쳐
# 그 안에 학습 문제를 네모 박스로 한 번만 보여준다.
_ANNOUNCE_LABELS = {"학습 문제 안내하기"}

# 활동1/활동2 행은 teacher 셀 안에서 "T : "로 시작하는 마지막 발문 줄을 제외한
# 번호 절차(1. 2. 3. ...) 부분만 테두리 박스로 감싼다.
_PROCEDURE_LABELS = {"활동1", "활동2"}

# 워드/한글의 한글 줄바꿈은 공백이 아니라 아무 글자 사이에서나 끊을 수 있어("전시학" /
# "습 상기"처럼) 좁은 열에서 단어 중간이 잘리는 경우가 있다. 학습내용 라벨은 정해진
# 문자열 몇 개뿐이라 자연스러운 줄바꿈에 맡기지 않고 원하는 지점에 직접 줄바꿈을 넣는다.
_CONTENT_LABEL_BREAKS = {
    "전시학습 상기": "전시학습\n상기",
    "학습 문제 안내하기": "학습 문제\n안내하기",
    "학습 내용 정리": "학습 내용\n정리",
}

# 가시성을 위해 교사·학생·유의점처럼 내용이 많은 열은 넓게 잡는다. 학습단계·시간은
# 기존 대비 1.5배로 넓혔고, 도구/유의점은 원래 폭(0.7in)의 1.5배(1.05in)로 넓혔다. 학습
# 내용은 원래 폭(0.69in)의 1.2배(0.83in)로 넓힌다 — 페이지 여백을 0.5in씩으로 줄여
# 본문 너비가 7.5in로 늘어난 덕분에 다른 열을 줄이지 않고도 여유 안에서 넓힐 수 있다.
_STAGE_COL_WIDTHS = [Inches(0.63), Inches(0.83), Inches(1.825), Inches(1.825), Inches(0.48), Inches(1.05)]
# 메타 표(차시/학습목표 등)와 학습단계 표의 가로 폭을 맞추기 위해, 메타 표 두 열의 합도
# 학습단계 표 전체 폭(6.64in)과 같게 비율을 유지해서 스케일한다.
_META_COL_WIDTHS = [Inches(1.552), Inches(5.088)]

# 평가 기준(상/중/하) 표 각 행의 최소 높이. 0.8in는 빈 여백이 너무 많아 0.55in로 줄인다
# (0.42in는 눌려 보였던 값이라 그 사이 절충값).
_EVAL_ROW_HEIGHT = Inches(0.55)

# 중첩 표(학습 문제 박스, 활동1/2 절차 박스)를 부모 셀 폭 그대로 채우면 셀 안쪽 여백
# 공간이 없어 테두리가 잘려 보인다. 이만큼 안쪽으로 줄여서 여유를 남긴다.
_BOX_WIDTH_INSET = Inches(0.2)

# 활동지는 초등학생이 직접 손으로 쓰는 실물 자료라 교수학습과정안 표(10pt)보다 글자를
# 크게 잡고, 문항마다 답을 쓸 수 있는 빈 줄을 남긴다. 한 페이지 안에 다 들어가야 하므로
# 줄 수는 1개로 제한한다(넉넉하되 넘치지 않게).
_WORKSHEET_FONT_SIZE = Pt(13)
_WORKSHEET_TITLE_FONT_SIZE = Pt(22)
_WORKSHEET_WRITING_LINES_PER_ITEM = 1

# 문서 전체(제목 두 줄 제외) 글자 크기를 10pt로 통일하되, 도구 및 자료/유의점 열은
# 항목이 많아 9pt로 그 열만 따로 줄인다.
_BODY_FONT_SIZE = Pt(10)
_TOOLS_NOTES_FONT_SIZE = Pt(9)


def _apply_column_widths(table: Table, widths: list[Inches]) -> None:
    """열 너비를 지정한다. autofit(내용에 맞춰 자동조정)이 켜져 있으면 워드/한컴오피스가
    지정한 너비를 무시하므로 끄고, tblGrid뿐 아니라 각 행의 셀 너비도 함께 맞춘다 —
    add_row()로 나중에 추가된 행은 그리드 너비만으로는 반영되지 않는 경우가 있다.
    가로 병합된 셀(학습 문제 박스, 학습 단계, 평가 기준 행)은 row.cells에서 같은 셀이
    여러 열 위치에 반복되므로, 병합된 만큼 너비를 합산해서 한 번만 지정한다."""
    table.autofit = False
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    for row in table.rows:
        seen_tc_ids: set[int] = set()
        for start_index, cell in enumerate(row.cells):
            tc_id = id(cell._tc)
            if tc_id in seen_tc_ids:
                continue
            seen_tc_ids.add(tc_id)
            span = 1
            while (
                start_index + span < len(row.cells)
                and id(row.cells[start_index + span]._tc) == tc_id
            ):
                span += 1
            cell.width = Emu(sum(int(w) for w in widths[start_index : start_index + span]))


def _shade_cell(cell: _Cell, hex_color: str) -> None:
    """표 셀에 배경색을 입힌다. 문단 배경색 XML 트릭보다 한컴오피스 호환성이 안정적이다."""
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color)
    cell_properties.append(shading)


def _shade_paragraph(paragraph, hex_color: str) -> None:
    """문단 한 줄에만 배경색을 입힌다(셀 전체가 아니라 라벨 줄 한 줄만 옅게 표시할 때 쓴다)."""
    paragraph_properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color)
    paragraph_properties.append(shading)


def _set_font_size(cell: _Cell, size: Pt) -> None:
    """셀 안 모든 문단의 글자 크기를 지정한다(문서 전체 기본값과 다르게 줄 때 사용)."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = size


def _set_run_font(run, name: str, size: Pt) -> None:
    """글꼴 이름과 크기를 함께 지정한다. 한글은 rFonts의 eastAsia 슬롯도 지정해야
    실제로 그 한글 폰트가 적용된다(ascii/hAnsi만 지정하면 한글에는 안 먹는다)."""
    run.font.size = size
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _set_worksheet_font(container, size: Pt | None = None) -> None:
    """활동지 전용 글꼴(나눔바른펜)을 셀 또는 문단 하나에 적용한다. size를 안 주면
    _WORKSHEET_FONT_SIZE를 쓴다(제목처럼 더 크게 쓸 땐 size를 따로 넘긴다)."""
    font_size = size or _WORKSHEET_FONT_SIZE
    paragraphs = container.paragraphs if hasattr(container, "paragraphs") else [container]
    for paragraph in paragraphs:
        for run in paragraph.runs:
            _set_run_font(run, _WORKSHEET_FONT_NAME, font_size)


def _no_wrap_cell(cell: _Cell) -> None:
    """셀 폭이 좁아 한글 라벨이 한 글자씩 세로로 줄바꿈되는 것을 막는다. 폭을 늘리는 대신
    한 줄을 유지하고 필요하면 옆으로 살짝 넘치도록 한다(학습 단계/학습 내용/시간, 메타 라벨처럼
    짧지만 좁은 열에 사용)."""
    cell_properties = cell._tc.get_or_add_tcPr()
    no_wrap = OxmlElement("w:noWrap")
    cell_properties.append(no_wrap)


def _style_header_cell(cell: _Cell) -> None:
    """표 헤더/라벨 셀(메타 표 라벨, 학습 단계 헤더 행, 도입/전개/정리)에 옅은 하늘색
    배경(#D4E0EC)과 검정 굵은 글씨를 적용하고 가운데 정렬한다. cell.text로 텍스트를
    먼저 채운 뒤 호출한다."""
    _shade_cell(cell, _HEADER_COLOR)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def _split_heading_body(text: str) -> tuple[str, str]:
    """LG-002 제목 줄 규칙(prompts.py)에 따라 teacher/student 필드의 첫 줄(◎/◘/♥ 제목)과
    그 다음 본문을 분리한다."""
    heading, _, body = text.partition("\n")
    return heading, body


def _add_bold_heading_cell(cell: _Cell, text: str) -> None:
    """제목 줄(첫 줄)만 굵게, 나머지는 일반 텍스트로 채운다. "활동1"/"활동2" 행의 student
    셀에 쓴다 — teacher 쪽 제목 줄이 굵게 나오는 것과 맞춘다."""
    heading, body = _split_heading_body(text)
    heading_paragraph = cell.paragraphs[0]
    heading_run = heading_paragraph.add_run(heading)
    heading_run.bold = True
    if body:
        cell.add_paragraph(body)


def _add_boxed_table(
    cell: _Cell, text: str, width: Emu, alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT
) -> None:
    """cell 안에 실선 테두리 박스(중첩 1x1 표)를 추가하고 text를 넣는다. width를 부모 셀
    폭 그대로 주면 셀 안쪽 여백이 없어 테두리가 잘려 보이므로 _BOX_WIDTH_INSET만큼 줄이고,
    중첩 표 자체를 가운데 정렬해 부모 셀 안에 여백을 두고 예쁘게 들어가도록 한다."""
    box_width = Emu(max(int(width) - int(_BOX_WIDTH_INSET), int(Inches(0.5))))
    box_table = cell.add_table(rows=1, cols=1)
    box_table.style = "Table Grid"
    box_table.autofit = False
    box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    box_table.columns[0].width = box_width
    box_cell = box_table.rows[0].cells[0]
    box_cell.width = box_width
    box_cell.text = text
    box_cell.paragraphs[0].alignment = alignment


def _add_announce_box(cell: _Cell, teacher_text: str, student_text: str, width: Emu) -> None:
    """"학습 문제 안내하기" 행의 교사·학생 병합 셀을 채운다: 제목 줄 두 개를 나란히
    굵게 쓴 뒤, 그 아래에 실선 테두리 박스로 학습 문제 문장을 병합 셀 폭 그대로 넣는다."""
    teacher_heading, body = _split_heading_body(teacher_text)
    student_heading, _ = _split_heading_body(student_text)

    heading_paragraph = cell.paragraphs[0]
    heading_run = heading_paragraph.add_run(f"{teacher_heading}    {student_heading}")
    heading_run.bold = True

    _add_boxed_table(cell, body, width, alignment=WD_ALIGN_PARAGRAPH.CENTER)


_NUMBERED_LINE = re.compile(r"^\d+\.")


def _add_procedure_box(cell: _Cell, teacher_text: str, width: Emu) -> None:
    """"활동1"/"활동2" 행의 teacher 셀을 채운다: 제목 줄은 굵게. 본문이 번호 절차
    (1. 2. 3. ...)면 게임·규칙 안내로 보고 테두리 박스로 감싸 눈에 띄게 하고, "-"로
    시작하는 서술형 문장(+ 중간중간 "T : " 발문 여러 개)이면 일반 활동이므로 박스 없이
    그대로 문단으로 나열한다(prompts.py 규칙 4)."""
    heading, body = _split_heading_body(teacher_text)
    body_lines = body.split("\n") if body else []

    heading_paragraph = cell.paragraphs[0]
    heading_run = heading_paragraph.add_run(heading)
    heading_run.bold = True

    if not body_lines:
        return

    is_numbered_procedure = bool(_NUMBERED_LINE.match(body_lines[0]))
    if is_numbered_procedure:
        _add_boxed_table(cell, body, width)
    else:
        for line in body_lines:
            cell.add_paragraph(line)


_S_LINE = re.compile(r"^S\d*\s*:\s")  # "S : "/"S1 : "/"S2 : " 등 학생 발화 줄 전체
_S_CONTINUATION = re.compile(r"^S[2-9]\d*\s*:\s")  # "S2 : "부터는 같은 응답 묶음의 연속


def _split_teacher_segments(body_lines: list[str]) -> list[list[str]]:
    """서술형 활동 본문을 "T : " 발문 하나씩으로 끝나는 묶음으로 나눈다(그 앞의
    "-" 설명 줄들 포함). 마지막에 T 없이 남는 줄이 있으면 그것도 한 묶음으로 둔다."""
    segments: list[list[str]] = []
    current: list[str] = []
    for line in body_lines:
        current.append(line)
        if line.startswith("T : "):
            segments.append(current)
            current = []
    if current:
        segments.append(current)
    return segments


def _split_student_groups(body_lines: list[str]) -> list[list[str]]:
    """학생 줄들을 교사의 각 "T : " 발문에 대응하는 묶음으로 나눈다. 한 묶음은
    [소개문 0줄 이상] + [응답(S/S1+S2+...) 1줄 이상]으로 구성된다 — 이미 응답을
    받은 묶음 뒤에 새 줄(소개문이든 "S :"/"S1 :" 새 응답이든)이 나오면 다음 발문에
    대응하는 새 묶음을 시작한다. "S2 : " 이후 번호는 같은 묶음의 추가 응답으로 본다."""
    groups: list[list[str]] = []
    current: list[str] = []
    has_response = False
    for line in body_lines:
        is_continuation = bool(_S_CONTINUATION.match(line))
        if current and has_response and not is_continuation:
            groups.append(current)
            current = []
            has_response = False
        current.append(line)
        if _S_LINE.match(line):
            has_response = True
    if current:
        groups.append(current)
    return groups


def _align_dialogue_bodies(
    teacher_lines: list[str], student_lines: list[str]
) -> tuple[list[str], list[str]]:
    """T:/S: 발화가 표에서 같은 줄 높이에 오도록, "T : " 하나로 끝나는 교사 묶음과
    그에 대응하는 학생 응답 묶음(S/S1+S2...) 쌍마다 짧은 쪽에 빈 줄을 끼워 넣는다.
    셀을 합치지 않고 각자 그대로 둔 채 줄 수만 맞춘다."""
    segments = _split_teacher_segments(teacher_lines)
    groups = _split_student_groups(student_lines)
    aligned_teacher: list[str] = []
    aligned_student: list[str] = []
    for i in range(max(len(segments), len(groups))):
        teacher_block = segments[i] if i < len(segments) else []
        student_block = groups[i] if i < len(groups) else []
        # 빈 줄을 내용 뒤가 아니라 앞에 붙인다 — 그래야 교사의 "T : " 줄과 학생의
        # 응답 줄처럼 각 묶음의 "핵심 줄"이 묶음 맨 아래에서 같은 높이로 만난다.
        aligned_teacher.extend([""] * max(0, len(student_block) - len(teacher_block)))
        aligned_teacher.extend(teacher_block)
        aligned_student.extend([""] * max(0, len(teacher_block) - len(student_block)))
        aligned_student.extend(student_block)
    return aligned_teacher, aligned_student


def _add_aligned_dialogue(
    teacher_cell: _Cell, student_cell: _Cell, teacher_text: str, student_text: str
) -> None:
    """"전시학습 상기"와 서술형(비게임) "활동1"/"활동2"에서 교사·학생 발화가 같은 줄
    높이에 나란히 오도록, 각 칸에 빈 줄을 끼워 넣는다(칸은 합치지 않는다)."""
    teacher_heading, teacher_body = _split_heading_body(teacher_text)
    student_heading, student_body = _split_heading_body(student_text)

    aligned_teacher, aligned_student = _align_dialogue_bodies(
        teacher_body.split("\n") if teacher_body else [],
        student_body.split("\n") if student_body else [],
    )

    for cell, heading, body_lines in (
        (teacher_cell, teacher_heading, aligned_teacher),
        (student_cell, student_heading, aligned_student),
    ):
        heading_run = cell.paragraphs[0].add_run(heading)
        heading_run.bold = True
        for line in body_lines:
            cell.add_paragraph(line)


def _render_worksheet_page(
    doc: Document, worksheet: Worksheet, worksheet_image: bytes | None = None
) -> None:
    """활동지 한 페이지를 doc에 그려 넣는다. 실제로 인쇄·복사해서 나눠주는 자료라
    배경색 채움 대신 흰 배경 + 테두리 박스(Table Grid) + 옅은 회색 라벨 줄로만
    구성한다(참고 이미지: 초등학교가 궁금해요/독서 기록장 — 색은 최소화, 테두리
    박스 위주). render_lesson_docx와 render_worksheet_docx가 함께 쓴다."""
    # 아이콘 없이 제목만 크게 — 활동지 전체에서 이모지를 아예 쓰지 않는다.
    worksheet_title = doc.add_heading(worksheet.title, level=1)
    worksheet_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_worksheet_font(worksheet_title, _WORKSHEET_TITLE_FONT_SIZE)
    for run in worksheet_title.runs:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # 반/번호/이름을 적을 칸.
    id_table = doc.add_table(rows=1, cols=3)
    id_table.style = "Table Grid"
    for cell, label in zip(id_table.rows[0].cells, ["반:", "번호:", "이름:"]):
        cell.text = label
        cell.paragraphs[0].runs[0].bold = True
    _set_worksheet_font(id_table.rows[0].cells[0])
    _set_worksheet_font(id_table.rows[0].cells[1])
    _set_worksheet_font(id_table.rows[0].cells[2])
    doc.add_paragraph()

    if worksheet_image is not None:
        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.add_run().add_picture(BytesIO(worksheet_image), width=Inches(4.5))

    mission_paragraph = doc.add_paragraph(worksheet.mission)
    mission_paragraph.runs[0].italic = True
    _set_worksheet_font(mission_paragraph)
    doc.add_paragraph()

    for section in worksheet.sections:
        # 섹션 전체를 테두리 박스(Table Grid) 하나로 감싼다 — 배경은 흰색 그대로 두고,
        # 라벨 줄만 옅은 회색으로 표시해 구분한다.
        section_table = doc.add_table(rows=1, cols=1)
        section_table.style = "Table Grid"
        section_cell = section_table.rows[0].cells[0]

        section_cell.paragraphs[0].text = section.step_label
        section_cell.paragraphs[0].runs[0].bold = True
        _shade_paragraph(section_cell.paragraphs[0], _WORKSHEET_LABEL_GRAY)
        section_cell.add_paragraph(section.instruction)

        if section.table is not None:
            visual_table = section_cell.add_table(
                rows=1 + len(section.table.rows), cols=len(section.table.headers)
            )
            visual_table.style = "Table Grid"
            for cell, header in zip(visual_table.rows[0].cells, section.table.headers):
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            for row_index, row_values in enumerate(section.table.rows, start=1):
                for cell, value in zip(visual_table.rows[row_index].cells, row_values):
                    cell.text = value
            for row in visual_table.rows:
                for cell in row.cells:
                    _set_worksheet_font(cell)

        for item_index, item in enumerate(section.items, start=1):
            # Word 내장 "List Number" 스타일은 한컴오피스에서 번호가 깨져 보이는
            # 경우가 있어, 스타일 대신 번호를 텍스트로 직접 써넣는다.
            section_cell.add_paragraph(f"{item_index}. {item}")
            # 초등학생이 직접 답을 쓸 수 있도록 문항마다 빈 줄을 남긴다(한 페이지 안에
            # 다 들어가야 하므로 한 줄만).
            for _ in range(_WORKSHEET_WRITING_LINES_PER_ITEM):
                section_cell.add_paragraph()

        _set_worksheet_font(section_cell)
        doc.add_paragraph()

    reflection_heading = doc.add_heading("생각해 보기", level=2)
    _set_worksheet_font(reflection_heading)
    reflection_table = doc.add_table(rows=1, cols=1)
    reflection_table.style = "Table Grid"
    reflection_cell = reflection_table.rows[0].cells[0]
    reflection_cell.paragraphs[0].text = (
        f"1. {worksheet.reflection_questions[0]}" if worksheet.reflection_questions else ""
    )
    for question_index, question in enumerate(worksheet.reflection_questions[1:], start=2):
        reflection_cell.add_paragraph(f"{question_index}. {question}")
    _set_worksheet_font(reflection_cell)


def render_worksheet_docx(worksheet: Worksheet, worksheet_image: bytes | None = None) -> bytes:
    """활동지 한 장만 담은 DOCX를 만든다. 전체 교수학습과정안 없이 활동지 디자인만
    빠르게 확인하거나, 활동지만 따로 인쇄하고 싶을 때 쓴다."""
    doc = Document()
    doc.styles["Normal"].font.size = _BODY_FONT_SIZE
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    _render_worksheet_page(doc, worksheet, worksheet_image)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_lesson_docx(lesson: LessonOutput, worksheet_image: bytes | None = None) -> bytes:
    """LessonOutput을 공식 수업지도안 양식 + 활동지 DOCX로 변환한다(개선-3).

    교사가 다운로드 후 워드/한글에서 직접 편집할 수 있도록 표 기반 문서로
    만든다. 디스크에 저장하지 않고 메모리(BytesIO)에서 bytes로 반환한다.
    실제 "다운로드" 버튼과 라우트는 main.py(D/E 소유)에서 연결해야 하며,
    이 함수는 그 라우트가 호출할 순수 변환 로직만 담당한다.

    worksheet_image: 활동지 대표 삽화(PNG/JPEG bytes). 있으면 활동지 제목
    바로 아래에 넣는다. 이미지 생성은 결제 연결이 필요해 선택 인자로 둔다
    (없으면 기존처럼 이모지 아이콘만으로 구성). lesson.worksheet가 애초에
    None이면(이 수업에 활동지가 필요 없다고 판단된 경우) 무시된다.
    """
    doc = Document()
    doc.styles["Normal"].font.size = _BODY_FONT_SIZE

    # 페이지 양쪽 여백을 워드/한글 기본값(1in)의 절반으로 줄인다.
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    heading = doc.add_heading("교수학습과정안", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_rows = [
        ("차시(시간)", lesson.lesson_time),
        ("교육 대상 - 학교급", lesson.school_level),
        ("교육 대상 - 학년", f"{lesson.grade}학년"),
        ("학습 주제", lesson.topic),
        ("관련 과목(영역)", lesson.subject),
        # 웹 UI(result.html)는 코드 옆에 성취기준 원문을 나란히 보여주는데, 여기서
        # achievement_code만 쓰면 DOCX로 내려받았을 때 한글 원문이 통째로 사라져 보인다.
        ("성취 기준", f"{lesson.achievement_code} {lesson.achievement_statement}"),
        ("학습 목표", "\n".join(f"- {objective}" for objective in lesson.learning_objectives)),
        # "활용 자료"만 다음 줄에 오도록 라벨 자체에 줄바꿈을 넣는다.
        ("학습 준비물 및\n활용 자료", ", ".join(lesson.materials)),
        ("AI·디지털 도구", lesson.ai_digital_tool),
    ]
    meta_table = doc.add_table(rows=len(meta_rows), cols=2)
    meta_table.style = "Table Grid"
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, (label, value) in enumerate(meta_rows):
        label_cell, value_cell = meta_table.rows[row_index].cells
        label_cell.text = label
        _style_header_cell(label_cell)
        value_cell.text = value
    _apply_column_widths(meta_table, _META_COL_WIDTHS)

    doc.add_paragraph()

    # 헤더는 2행이다 — 학습 단계/학습 내용/시간/도구자료는 두 행에 걸쳐 세로 병합하고,
    # "교수·학습활동"은 교사/학생 두 열에 걸쳐 가로로 병합한 뒤 그 아래(2행)에
    # "교사 활동"/"학생 활동" 하위 헤더를 둔다(실제 공식 양식과 동일한 2단 헤더).
    stage_table = doc.add_table(rows=2, cols=6)
    stage_table.style = "Table Grid"
    stage_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row0 = stage_table.rows[0].cells
    header_row1 = stage_table.rows[1].cells

    stage_header = header_row0[0].merge(header_row1[0])
    content_header = header_row0[1].merge(header_row1[1])
    activity_header = header_row0[2].merge(header_row0[3])
    time_header = header_row0[4].merge(header_row1[4])
    tools_header = header_row0[5].merge(header_row1[5])

    stage_header.text = "학습 단계"
    content_header.text = "학습 내용"
    activity_header.text = "교수 학습 활동"
    time_header.text = "시간\n(분)"
    tools_header.text = "도구 및 자료(□) / 유의점(◆)"
    header_row1[2].text = "교사 활동"
    header_row1[3].text = "학생 활동"

    # 학습 단계(0)/학습 내용(1)/시간(4)은 줄바꿈을 막아 가로 한 줄을 유지한다. 학습 내용은
    # _CONTENT_LABEL_BREAKS로 원하는 지점에 이미 명시적 줄바꿈을 넣어두므로, noWrap을
    # 걸어도 그 지점에서는 정상적으로 끊기고 그 외 지점에서 추가로 잘리지 않는다.
    for cell in [stage_header, content_header, activity_header, time_header, tools_header, header_row1[2], header_row1[3]]:
        _style_header_cell(cell)
    for cell in [stage_header, content_header, time_header]:
        _no_wrap_cell(cell)
    _set_font_size(tools_header, _TOOLS_NOTES_FONT_SIZE)

    stages = [
        ("도입", lesson.lesson_stages.intro),
        ("전개", lesson.lesson_stages.development),
        ("정리", lesson.lesson_stages.wrap_up),
    ]
    for label, activities in stages:
        first_row_index = len(stage_table.rows)
        for activity in activities:
            row = stage_table.add_row().cells
            row[1].text = _CONTENT_LABEL_BREAKS.get(activity.content_label, activity.content_label)
            _no_wrap_cell(row[1])
            if activity.content_label in _ANNOUNCE_LABELS:
                merged = row[2].merge(row[3])
                merged_width = Emu(int(_STAGE_COL_WIDTHS[2]) + int(_STAGE_COL_WIDTHS[3]))
                _add_announce_box(merged, activity.teacher, activity.student, merged_width)
            elif activity.content_label in _PROCEDURE_LABELS:
                _, teacher_body = _split_heading_body(activity.teacher)
                teacher_body_lines = teacher_body.split("\n") if teacher_body else []
                is_numbered = bool(teacher_body_lines) and _NUMBERED_LINE.match(
                    teacher_body_lines[0]
                )
                if is_numbered:
                    # 게임·규칙 형식: 기존처럼 교사 칸만 절차를 박스로, 학생 칸은 따로.
                    _add_procedure_box(row[2], activity.teacher, _STAGE_COL_WIDTHS[2])
                    _add_bold_heading_cell(row[3], activity.student)
                else:
                    # 서술형 형식: 칸은 합치지 않고, "T : "/"S : " 쌍이 같은 줄 높이에
                    # 오도록 짧은 쪽에 빈 줄만 끼워 넣는다.
                    _add_aligned_dialogue(row[2], row[3], activity.teacher, activity.student)
            elif activity.content_label == "전시학습 상기":
                # S1/S2처럼 학생 응답이 여러 줄이면 교사 쪽과 줄 수가 달라지므로
                # 여기도 같은 정렬 규칙을 적용한다.
                _add_aligned_dialogue(row[2], row[3], activity.teacher, activity.student)
            else:
                row[2].text = activity.teacher
                row[3].text = activity.student
            row[4].text = activity.time
            _no_wrap_cell(row[4])
            row[5].text = "\n".join(
                [f"□ {tool}" for tool in activity.tools]
                + [f"◆ {note}" for note in activity.notes]
            )
            _set_font_size(row[5], _TOOLS_NOTES_FONT_SIZE)
        # 학습 단계 열은 그 단계의 모든 학습내용 행에 걸쳐 세로로 병합한다.
        last_row_index = len(stage_table.rows) - 1
        stage_cell = stage_table.rows[first_row_index].cells[0]
        if last_row_index > first_row_index:
            stage_cell = stage_cell.merge(stage_table.rows[last_row_index].cells[0])
        stage_cell.text = label
        _style_header_cell(stage_cell)
        _no_wrap_cell(stage_cell)

    eval_header_row = stage_table.add_row().cells
    eval_header_cell = eval_header_row[0]
    for col_index in range(1, 6):
        eval_header_cell = eval_header_cell.merge(eval_header_row[col_index])
    eval_header_cell.text = "평가 기준"
    _style_header_cell(eval_header_cell)

    eval_body_row = stage_table.add_row().cells
    eval_body_cell = eval_body_row[0]
    for col_index in range(1, 6):
        eval_body_cell = eval_body_cell.merge(eval_body_row[col_index])
    eval_table = eval_body_cell.add_table(rows=3, cols=2)
    eval_table.style = "Table Grid"
    eval_table.autofit = False
    eval_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    eval_level_width = Inches(0.5)
    # 병합된 셀 전체 폭(6.64in)까지 늘리면 설명 한 줄이 페이지 끝까지 늘어져 오히려
    # 읽기 힘들어지지만(4.2in 시도), 너무 좁으면 표 좌우 여백이 어색해 보인다. 그
    # 중간값으로 넓혀 가운데 정렬 여백을 자연스럽게 남긴다.
    eval_desc_width = Inches(5.5)
    eval_table.columns[0].width = eval_level_width
    eval_table.columns[1].width = eval_desc_width
    for row, (level, description) in zip(
        eval_table.rows,
        [
            ("상", lesson.evaluation_criteria.high),
            ("중", lesson.evaluation_criteria.mid),
            ("하", lesson.evaluation_criteria.low),
        ],
    ):
        row.height = _EVAL_ROW_HEIGHT
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        level_cell, desc_cell = row.cells
        level_cell.width = eval_level_width
        desc_cell.width = eval_desc_width
        level_cell.text = level
        level_cell.paragraphs[0].runs[0].bold = True
        level_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc_cell.text = description

    _apply_column_widths(stage_table, _STAGE_COL_WIDTHS)

    # 활동지는 이 수업에 실제로 필요한 경우에만 생성된다(prompts.py 규칙 7). 없으면
    # (worksheet is None) 이 페이지 전체를 건너뛴다.
    if lesson.worksheet is not None:
        doc.add_page_break()  # 별도 페이지로 분리(교사가 이 페이지만 따로 인쇄할 수 있도록)
        _render_worksheet_page(doc, lesson.worksheet, worksheet_image)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
